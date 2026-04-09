"""
Module for extracting text from PDF and DOCX files
"""
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import io

def extract_text_from_file(file_obj, file_type):
    """
    Extract text from PDF or DOCX file
    
    Args:
        file_obj: File object (io.BytesIO)
        file_type: File type ('pdf', 'docx', or 'doc')
    
    Returns:
        str: Extracted text
    """
    try:
        if file_type.lower() == 'pdf':
            return extract_text_from_pdf(file_obj)
        elif file_type.lower() in ['docx', 'doc']:
            return extract_text_from_docx(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        raise Exception(f"Error extracting text from file: {str(e)}")

def extract_text_from_pdf(file_obj):
    """
    Extract text from PDF file
    
    Args:
        file_obj: File object (io.BytesIO)
    
    Returns:
        str: Extracted text
    """
    try:
        # Reset file pointer to beginning
        file_obj.seek(0)
        text = pdf_extract_text(file_obj)
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file_obj):
    """
    Extract text from DOCX file
    
    Args:
        file_obj: File object (io.BytesIO)
    
    Returns:
        str: Extracted text
    """
    try:
        # Reset file pointer to beginning
        file_obj.seek(0)
        doc = Document(file_obj)
        
        # Extract text from all paragraphs
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += '\n' + cell.text
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")
